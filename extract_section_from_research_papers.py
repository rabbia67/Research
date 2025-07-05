"""
extract_intros.py  ── one‑click version with debug
Dependencies:
    pip install pymupdf python-docx tqdm
"""

import pathlib, re, fitz  # PyMuPDF is imported as fitz
from docx import Document
from tqdm import tqdm
from docx.enum.text import WD_BREAK
from docx.enum.style import WD_STYLE_TYPE

# ── 1) EDIT THESE TWO LINES ONLY ──────────────────────────────────────────────
PAPERS_DIR = pathlib.Path(r"C:\Rabbia MS data\Thesis\thesis").expanduser().resolve()
OUT_DOCX = pathlib.Path(r"C:\Users\Rabbia\Desktop\project2\CV\introductions.docx").expanduser().resolve()
# ──────────────────────────────────────────────────────────────────────────────

# More flexible "Introduction" section matching
INTRO_PAT = re.compile(r"^\s*(\d{0,2}\.?\s*)?(introduction)\b", re.I | re.M)

# Stop at the next major heading or known section title
STOP_PAT = re.compile(r"^\s*(\d+\.?)\s+[A-Z].+$", re.M)
KNOWN_TITLES = (
    "related work", "literature review", "background", "methods",
    "materials", "methodology", "experimental", "data",
    "proposed", "approach", "results", "discussion", "analysis"
)
STOP_UNKNOWN = re.compile(rf"^\s*(\d+\.?)?\s*({'|'.join(KNOWN_TITLES)})\b", re.I | re.M)

def _clean(t: str) -> str:
    return re.sub(r"[ \t]+\n", "\n", t).strip()

def intro_from_text(text: str) -> str | None:
    text = _clean(text)
    m = INTRO_PAT.search(text)
    if not m:
        return None
    stop = STOP_PAT.search(text, m.end()) or STOP_UNKNOWN.search(text, m.end())
    end = stop.start() if stop else len(text)
    intro = _clean(text[m.start():end])
    return intro if len(intro.split()) > 60 else None

def intro_from_pdf(p: pathlib.Path) -> str | None:
    doc = fitz.open(p)
    try:
        text = "\n".join(pg.get_text("text") for pg in doc)
    finally:
        doc.close()
    # Uncomment to debug text extraction:
    # print(f"\n--- {p.name} content start ---\n{text[:1000]}\n--- end ---\n")
    return intro_from_text(text)

def intro_from_docx(p: pathlib.Path) -> str | None:
    from docx import Document as _D
    doc = _D(p)
    full_text = "\n".join(par.text.strip() for par in doc.paragraphs if par.text.strip())
    # Uncomment to debug text extraction:
    # print(f"\n--- {p.name} content start ---\n{full_text[:1000]}\n--- end ---\n")
    return intro_from_text(full_text)

def extract_all():
    out = Document()

    if 'Heading 1' not in out.styles:
        out.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)

    files = sorted(PAPERS_DIR.glob("*"))
    if not files:
        print(f"⚠ No files found in {PAPERS_DIR}")
        return

    found_any = False

    for file in tqdm(files, desc="Extracting introductions"):
        print(f"→ Processing: {file.name}")
        intro = None

        if file.suffix.lower() == ".pdf":
            try:
                intro = intro_from_pdf(file)
            except Exception as e:
                print(f"✗ Error reading PDF {file.name}: {e}")
                continue
        elif file.suffix.lower() == ".docx":
            try:
                intro = intro_from_docx(file)
            except Exception as e:
                print(f"✗ Error reading DOCX {file.name}: {e}")
                continue
        else:
            print(f"Skipping unsupported file type: {file.name}")
            continue

        if intro:
            out.add_heading(file.stem, level=1)
            out.add_paragraph(intro)
            out.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)
            print(f"✓ Intro extracted from {file.name}")
            found_any = True
        else:
            print(f"⚠ No intro found in {file.name}")

    if found_any:
        out.save(OUT_DOCX)
        print(f"\n✅ Saved introductions to: {OUT_DOCX}")
    else:
        print("\n⚠ No introductions were extracted. Please check your documents or regex.")

if __name__ == "__main__":
    extract_all()
